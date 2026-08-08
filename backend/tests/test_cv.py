from io import BytesIO
from unittest.mock import MagicMock, patch

import docx


def _create_docx_bytes() -> bytes:
    document = docx.Document()
    document.add_heading("John Doe", 0)
    document.add_paragraph("Software Engineer with 5 years experience.")
    document.add_paragraph("Skills: Python, FastAPI, React")
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


MOCK_CV_PROFILE = {
    "name": "John Doe",
    "skills": ["Python", "FastAPI"],
    "experience": [{"title": "Engineer", "company": "Tech Co", "duration": "3 years"}],
    "education": [{"degree": "BS CS", "institution": "University", "year": "2020"}],
    "projects": [{"name": "Lead Gen App", "description": "SaaS platform"}],
    "services": ["Web Development", "API Design"],
    "tools": ["Git", "Docker"],
    "technologies": ["Python", "FastAPI"],
}

MOCK_SUMMARIES = {
    "professional_summary": "Experienced software engineer.",
    "skills_summary": "Strong Python and FastAPI skills.",
    "services_summary": "Offers web development services.",
    "experience_summary": "5 years in software engineering.",
}


@patch("app.services.cv_service.GroqService")
def test_upload_cv(mock_groq_class, client, groq_auth_headers):
    mock_groq = MagicMock()
    mock_groq.extract_cv_profile.return_value = MOCK_CV_PROFILE
    mock_groq.generate_summaries.return_value = MOCK_SUMMARIES
    mock_groq_class.return_value = mock_groq

    docx_bytes = _create_docx_bytes()
    response = client.post(
        "/api/cv/upload",
        headers=groq_auth_headers,
        files={"file": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 201
    data = response.json()
    assert data["message"] == "CV uploaded and processed successfully"
    assert data["profile"]["name"] == "John Doe"


@patch("app.services.cv_service.GroqService")
def test_get_cv_profile(mock_groq_class, client, groq_auth_headers):
    mock_groq = MagicMock()
    mock_groq.extract_cv_profile.return_value = MOCK_CV_PROFILE
    mock_groq.generate_summaries.return_value = MOCK_SUMMARIES
    mock_groq_class.return_value = mock_groq

    docx_bytes = _create_docx_bytes()
    client.post(
        "/api/cv/upload",
        headers=groq_auth_headers,
        files={"file": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    response = client.get("/api/cv/profile", headers=groq_auth_headers)
    assert response.status_code == 200
    assert response.json()["name"] == "John Doe"


@patch("app.services.cv_service.GroqService")
def test_get_cv_raw(mock_groq_class, client, groq_auth_headers):
    mock_groq = MagicMock()
    mock_groq.extract_cv_profile.return_value = MOCK_CV_PROFILE
    mock_groq.generate_summaries.return_value = MOCK_SUMMARIES
    mock_groq_class.return_value = mock_groq

    docx_bytes = _create_docx_bytes()
    client.post(
        "/api/cv/upload",
        headers=groq_auth_headers,
        files={"file": ("resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    response = client.get("/api/cv/raw", headers=groq_auth_headers)
    assert response.status_code == 200
    assert "John Doe" in response.json()["raw_text"]


def test_cv_profile_not_found(client, auth_headers):
    response = client.get("/api/cv/profile", headers=auth_headers)
    assert response.status_code == 200
    assert response.json() is None
